"""
Dokuman okuyucu modulu.
Farkli dosya formatlarindan duz metin cikarir.
Desteklenen: .txt, .md, .pdf, .docx
"""
import re
from pathlib import Path
# Desteklenen uzantilar (beyaz liste)
DESTEKLENEN_UZANTILAR = {".txt", ".md", ".pdf", ".docx"}
# Maksimum dosya boyutu (bayt)
MAKS_DOSYA_BOYUTU = 10 * 1024 * 1024        # 10 MB
# Taranmis PDF tespiti icin sayfa basina minimum karakter
SAYFA_BASINA_MIN_KARAKTER = 50
class OkumaHatasi(Exception):
    """Dosya okunamadiginda firlatilir."""
    pass
def dosya_adini_temizle(ad: str) -> str:
    """
    Dosya adindan tehlikeli karakterleri temizler.
    Dizin gecisi (path traversal) saldirilarini engeller.
    """
    # Sadece dosya adini al, dizin bilesenlerini at
    ad = Path(ad).name
    # Alfanumerik, bosluk, nokta, tire ve alt cizgi disindakileri degistir
    temiz = re.sub(r"[^\w\s.-]", "_", ad)
    # Ardisik noktalari tekile indir (..\.. gibi kaliplari kir)
    temiz = re.sub(r"\.{2,}", ".", temiz)
    # Bosluklari alt cizgiye cevir
    temiz = temiz.strip().replace(" ", "_")
    return temiz or "adsiz_dosya"
def uzanti_gecerli_mi(dosya_adi: str) -> bool:
    """Uzanti beyaz listede mi?"""
    return Path(dosya_adi).suffix.lower() in DESTEKLENEN_UZANTILAR
def metni_temizle(metin: str) -> str:
    """
    Cikarilmis metni normallestirir.
    Fazla bosluk ve bos satirlari temizler.
    """
    # Windows satir sonlarini normallestir
    metin = metin.replace("\r\n", "\n").replace("\r", "\n")
    # Ucten fazla ardisik bos satiri ikiye indir
    metin = re.sub(r"\n{3,}", "\n\n", metin)
    # Satir sonundaki bosluklari temizle
    satirlar = [satir.rstrip() for satir in metin.split("\n")]
    return "\n".join(satirlar).strip()
# ==========================================================
# FORMAT OZEL OKUYUCULAR
# ==========================================================
def duz_metin_oku(yol: Path) -> str:
    """.txt ve .md dosyalarini okur."""
    for kodlama in ("utf-8", "utf-8-sig", "latin-1", "cp1254"):
        try:
            return yol.read_text(encoding=kodlama)
        except UnicodeDecodeError:
            continue
    raise OkumaHatasi("Dosya kodlamasi cozulemedi")
def pdf_oku(yol: Path) -> str:
    """
    PDF dosyasindan metin cikarir.
    Taranmis (goruntu tabanli) PDF'lerde metin bulunamaz;
    bu durumda anlamli bir hata mesaji uretilir.
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        raise OkumaHatasi(
            "pypdf kutuphanesi kurulu degil. "
            "Kurulum: pip install pypdf"
        )
    try:
        reader = PdfReader(str(yol))
    except Exception as hata:
        raise OkumaHatasi(f"PDF acilamadi: {hata}")
    if reader.is_encrypted:
        raise OkumaHatasi("PDF sifreli. Sifresiz bir kopya yukleyin.")
    sayfa_metinleri = []
    for numara, sayfa in enumerate(reader.pages, start=1):
        try:
            metin = sayfa.extract_text() or ""
        except Exception:
            metin = ""
        if metin.strip():
            sayfa_metinleri.append(metin.strip())
    if not sayfa_metinleri:
        raise OkumaHatasi(
            "PDF'den metin cikarilamadi. Dosya taranmis (goruntu tabanli) "
            "olabilir. Bu durumda OCR islemi gerekir."
        )
    tam_metin = "\n\n".join(sayfa_metinleri)
    # Taranmis PDF sezgisel tespiti
    ortalama = len(tam_metin) / len(reader.pages)
    if ortalama < SAYFA_BASINA_MIN_KARAKTER:
        raise OkumaHatasi(
            f"PDF'den cok az metin cikarildi (sayfa basina {ortalama:.0f} "
            f"karakter). Dosya buyuk olasilikla taranmis bir belgedir."
        )
    return tam_metin
def docx_oku(yol: Path) -> str:
    """
    Word (.docx) dosyasindan metin cikarir.
    Paragraflarin yani sira tablo hucrelerini de okur.
    """
    try:
        from docx import Document
    except ImportError:
        raise OkumaHatasi(
            "python-docx kutuphanesi kurulu degil. "
            "Kurulum: pip install python-docx"
        )
    try:
        belge = Document(str(yol))
    except Exception as hata:
        raise OkumaHatasi(
            f"Word dosyasi acilamadi: {hata}. "
            "Dosyanin .docx formatinda oldugundan emin olun "
            "(eski .doc formati desteklenmiyor)."
        )
    parcalar = []
    # Paragraflar
    for paragraf in belge.paragraphs:
        metin = paragraf.text.strip()
        if not metin:
            continue
        # Baslik stillerini markdown basligina cevir
        # paragraf.style bazi belgelerde None olabilir
        stil_nesnesi = getattr(paragraf, "style", None)
        stil = (getattr(stil_nesnesi, "name", None) or "").lower()
        if stil.startswith("heading 1") or stil == "title":
            parcalar.append(f"# {metin}")
        elif stil.startswith("heading 2"):
            parcalar.append(f"## {metin}")
        elif stil.startswith("heading"):
            parcalar.append(f"### {metin}")
        else:
            parcalar.append(metin)
    # Tablolar - hucreler bosluk ile ayrilarak duz metne cevrilir
    for tablo in belge.tables:
        for satir in tablo.rows:
            hucreler = [h.text.strip() for h in satir.cells if h.text.strip()]
            if hucreler:
                parcalar.append(" | ".join(hucreler))
    if not parcalar:
        raise OkumaHatasi("Word dosyasinda metin bulunamadi.")
    return "\n\n".join(parcalar)
# ==========================================================
# ANA GIRIS NOKTASI
# ==========================================================
def metni_cikar(yol) -> str:
    """
    Dosya uzantisina gore uygun okuyucuyu secer ve metni dondurur.
    Raises:
        OkumaHatasi: dosya okunamazsa veya desteklenmiyorsa
    """
    yol = Path(yol)
    if not yol.exists():
        raise OkumaHatasi(f"Dosya bulunamadi: {yol}")
    if not yol.is_file():
        raise OkumaHatasi(f"Bu bir dosya degil: {yol}")
    boyut = yol.stat().st_size
    if boyut == 0:
        raise OkumaHatasi("Dosya bos")
    if boyut > MAKS_DOSYA_BOYUTU:
        raise OkumaHatasi(
            f"Dosya cok buyuk ({boyut/1024/1024:.1f} MB). "
            f"Sinir: {MAKS_DOSYA_BOYUTU/1024/1024:.0f} MB"
        )
    uzanti = yol.suffix.lower()
    if uzanti not in DESTEKLENEN_UZANTILAR:
        raise OkumaHatasi(
            f"Desteklenmeyen format: {uzanti}. "
            f"Desteklenenler: {', '.join(sorted(DESTEKLENEN_UZANTILAR))}"
        )
    if uzanti in (".txt", ".md"):
        ham = duz_metin_oku(yol)
    elif uzanti == ".pdf":
        ham = pdf_oku(yol)
    elif uzanti == ".docx":
        ham = docx_oku(yol)
    else:
        raise OkumaHatasi(f"Okuyucu tanimlanmamis: {uzanti}")
    temiz = metni_temizle(ham)
    if len(temiz) < 20:
        raise OkumaHatasi(
            f"Cikarilan metin cok kisa ({len(temiz)} karakter). "
            "Dosya icerigi bos veya okunamiyor olabilir."
        )
    return temiz
def dosya_bilgisi(yol) -> dict:
    """Bir dosya hakkinda ozet bilgi dondurur (okumadan)."""
    yol = Path(yol)
    return {
        "ad": yol.name,
        "uzanti": yol.suffix.lower(),
        "boyut": yol.stat().st_size if yol.exists() else 0,
        "destekleniyor": uzanti_gecerli_mi(yol.name),
    }
